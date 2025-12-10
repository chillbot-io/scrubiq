"""Word document (.docx) extraction."""

from pathlib import Path

from .base import Extractor, ExtractionError

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocxExtractor(Extractor):
    """Extract text from Word documents."""

    @property
    def extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, path: Path) -> str:
        """Extract text from paragraphs and tables."""
        if not HAS_DOCX:
            raise ExtractionError("python-docx not installed. Run: pip install python-docx")

        try:
            doc = Document(path)
            text_parts = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            # Extract from headers and footers
            for section in doc.sections:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)

                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)

            return "\n".join(text_parts)

        except Exception as e:
            raise ExtractionError(f"Failed to extract from {path}: {e}")
